#!/usr/bin/env python3
"""Generate CONNECT_INTEL_CRM_PRD.docx — detailed CRM code & flow PRD."""

from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "docs" / "CONNECT_INTEL_CRM_PRD.docx"


def set_doc_defaults(doc):
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)
    style.element.rPr.rFonts.set(qn("w:eastAsia"), "Calibri")


def add_title_page(doc):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Connect Intel")
    run.bold = True
    run.font.size = Pt(28)

    p2 = doc.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run("CRM Platform — Technical PRD & Code Flow")
    r2.font.size = Pt(18)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run("\nVersion: July 2026\n").font.size = Pt(12)
    meta.add_run("Live: https://connectintel.net\n").font.size = Pt(12)
    meta.add_run("Confidential — Internal & partner review\n").font.size = Pt(11)

    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text):
    doc.add_paragraph(text)


def bullets(doc, items):
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def table(doc, headers, rows):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = "Table Grid"
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
    for ri, row in enumerate(rows):
        cells = t.rows[ri + 1].cells
        for ci, val in enumerate(row):
            cells[ci].text = str(val)
    doc.add_paragraph()


def build(doc):
    add_title_page(doc)

    # 1
    h1(doc, "1. Executive Summary")
    para(doc, (
        "Connect Intel is a B2B revenue workspace built for exporters, manufacturers, distributors, "
        "and trading teams selling across borders. The product is a single-page application (SPA) that "
        "unifies CRM pipeline management, marketing automation, team collaboration, and AI-assisted "
        "sales intelligence — without full page reloads."
    ))
    para(doc, (
        "The CRM is lead-centric: every sales motion revolves around savedLeads (pipeline entries). "
        "Each lead wraps a contact record plus a rich crm object (status, tasks, meetings, deals, "
        "emails, activities, notes). Companies, contacts, deals, and campaigns are promoted to "
        "first-class views while preserving backward compatibility with the pipeline hub model."
    ))
    bullets(doc, [
        "Primary stack: Vite + React 19 frontend, Node 22 Vercel serverless API, hybrid JSON store + PostgreSQL hot paths, optional Meilisearch.",
        "Multi-tenant: company organizations with RBAC (org admin, manager, rep) plus marketing roles.",
        "Production: connectintel.net on Vercel; data on Supabase when deployed.",
        "Design north star: HubSpot-grade objects, timeline, permissions, and reporting — original UI only.",
    ])

    # 2
    h1(doc, "2. Product Vision & Modules")
    table(doc, ["Module", "User value", "Primary UI"], [
        ("CRM / Pipeline", "Kanban/table/deals views; lead workspace; sequences", "PipelinePanel, LeadWorkspace"),
        ("Marketing", "Campaigns, audiences, templates, automations, analytics", "MarketingPanel"),
        ("Home / Analytics", "Dashboard, My Day, team KPIs, activity log", "OverviewPanel, TeamDashboardPanel"),
        ("Team / Workspace", "Invites, permissions, hierarchy, email domain", "TeamPanel"),
        ("AI Copilot", "CRM queries, research, outreach drafts, navigation", "ConnectAssistant"),
        ("Platform operator", "Customer support, imports, system status (staff only)", "PlatformAdminHome"),
    ])
    para(doc, (
        "Reference: docs/CRM_PLATFORM_BLUEPRINT.md is the master evolution document. "
        "This PRD describes how the live codebase implements that vision today."
    ))

    # 3
    h1(doc, "3. Technology Stack")
    h2(doc, "3.1 Frontend")
    bullets(doc, [
        "Vite 8 + React 19 + Tailwind 4 — frontend/ builds to site/",
        "SPA routing via AppShell + appHistory.js (URL query params, not React Router pages)",
        "Global command palette: CommandPalette.jsx (⌘K / Ctrl+K)",
        "Design tokens: platform-design-system.css",
        "PWA: service worker precache for offline shell",
    ])
    h2(doc, "3.2 Backend API")
    bullets(doc, [
        "Single Vercel serverless entry: api/index.js",
        "~130 routes via lib/server/apiRouteRegistry.js",
        "Handlers: lib/server/handlers/*.js — CORS, requireUser(), JSON responses",
        "Optional workers: Railway BullMQ for heavy email (workers/index.mjs)",
        "Observability: Sentry on API routes; Grafana metrics cron",
    ])
    h2(doc, "3.3 Data layer")
    table(doc, ["Layer", "Technology", "Purpose"], [
        ("JSON store", "Supabase collections or local SQLite", "Users, orgs, full lead blobs, marketing metadata"),
        ("SQL hot path", "PostgreSQL (Supabase)", "pipeline_leads, activities, deals, email queue"),
        ("Search", "Meilisearch (optional)", "Platform search ⌘K, pipeline index"),
        ("Cache", "In-memory + Redis optional", "Session user, dashboard snapshots"),
    ])

    # 4
    h1(doc, "4. User Types & Access Model")
    h2(doc, "4.1 Platform operator (Connect Intel staff)")
    para(doc, (
        "Identified by isPlatformAdmin on the session user. Server checks: user.role === 'admin', "
        "email in ADMIN_EMAILS env, or invite@connectintel.net (COMPANY_MAILBOX). "
        "Operators see buildOperatorNavSections() — admin-home, admin-customers, admin imports, integrations."
    ))
    para(doc, (
        "Guards (three layers): frontend/src/lib/platformOperator.js routing; PlatformOperatorGate.jsx; "
        "API requireAdmin() in lib/server/auth.js. Customers must never see operator panels or infra error text."
    ))
    h2(doc, "4.2 Company organization user")
    bullets(doc, [
        "accountType: 'company', organizationId set",
        "Org admin (isOrgAdmin): full org pipeline, team settings, permissions, imports",
        "Manager (pipelineRole: 'manager'): team-scoped pipeline via pipelineManagerScope.js",
        "Rep (pipelineRole: 'member'): own + collaborated + unassigned pool leads",
        "Marketing roles: marketing_manager, marketing_executive, marketing_readonly",
    ])
    h2(doc, "4.3 Individual seller")
    para(doc, (
        "accountType: 'individual', no organizationId. Own savedLeads where userId matches. "
        "Full admin-equivalent permissions on personal workspace."
    ))

    # 5
    h1(doc, "5. Authentication & Session Flow")
    h3(doc, "5.1 Sign-in paths")
    table(doc, ["Method", "Entry", "Server flow"], [
        ("Google OAuth", "AuthPage / GoogleSignIn", "auth-session.js POST → verifyGoogleCredential → upsertUser"),
        ("Email + password", "AuthPage tabs", "emailPasswordAuth.js"),
        ("Demo", "Dev only", "verifyDemoProfile → upsertUser"),
    ])
    h3(doc, "5.2 Session lifecycle")
    bullets(doc, [
        "JWT signed in sessionJwt.js; cookie connect_intel_session or Bearer token",
        "GET /api/auth/session refreshes from database (refreshSessionFromDatabase) — live org + ADMIN_EMAILS",
        "buildOrgUserResponse() in organizations.js is canonical session shape",
        "Login reads AUTH_STORE_COLLECTIONS slice only — fast auth without full CRM load",
    ])
    h3(doc, "5.3 Post-auth routing")
    bullets(doc, [
        "Incomplete onboarding → OnboardingModal (individual vs company)",
        "POST /api/onboarding/complete → completeOnboarding() creates org + membership",
        "Platform operators: onboarding auto-complete; land on admin-home",
        "Customers: default overview/pipeline; operator URLs sanitized",
    ])

    # 6
    h1(doc, "6. Multi-Tenant Organization Model")
    h2(doc, "6.1 Core entities (JSON store)")
    bullets(doc, [
        "organizations[] — name, domain, planTier, crmSettings, emailDomain, workspacePreset",
        "organizationMemberships[] — userId, organizationId, role, pipelineRole, marketingRole, canSearch",
        "organizationInvites[] — pending invites with expiry",
    ])
    h2(doc, "6.2 Tenant isolation")
    bullets(doc, [
        "organizationId on all CRM writes; tenantIsolation.js sanitizes cross-actor data",
        "repPipelineEntryVisible() — assignee, saver, collaborators, calendar participants, unassigned",
        "pipelineOwnerUserId() — assignee if set, else savedByUserId / userId",
        "canAssignLead() — managers/admins任意; reps claim unassigned or reassign own",
    ])
    h2(doc, "6.3 RBAC permission matrix")
    para(doc, (
        "Nine org actions enforced via assertOrgPermission() in permissionEnforce.js: "
        "view_all_leads, edit_leads, delete_leads, export_leads, view_analytics, access_marketing, "
        "send_campaigns, manage_team, manage_billing. Matrix roles: admin, manager, rep, marketing_*."
    ))
    h2(doc, "6.4 SQL sync (enterprise path)")
    bullets(doc, [
        "deferOrgSqlSync / deferMemberSqlSync — JSON → Supabase profiles + organizations",
        "USE_PIPELINE_LEADS_TABLE — hot pipeline reads from SQL not full JSON shard",
        "USE_PIPELINE_HIERARCHY_RBAC — rep/manager/admin SQL scoping",
        "Scripts: pipeline:sync, activities:backfill, meili:sync, org:sql-backfill",
    ])

    # 7
    h1(doc, "7. Core Data Objects")
    h2(doc, "7.1 Saved lead (pipeline hub)")
    para(doc, (
        "Storage: savedLeads[] — each entry has lead (contact fields), crm (nested CRM), "
        "organizationId, userId, savedByUserId, assignedToUserId, contactId, companyId."
    ))
    table(doc, ["CRM field", "Contents"], [
        ("status", "Pipeline stage: new, contacted, follow_up, replied, won, active_trading, lost"),
        ("activities[]", "Logged CRM actions with actor_id"),
        ("tasks[] / meetings[]", "Follow-ups and calendar items"),
        ("deals[]", "Nested deals per lead (amount, stage, close date)"),
        ("emails[]", "Sent/received trail (extension + Resend)"),
        ("notes, tagIds, leadScore", "Collaboration and scoring"),
    ])
    h2(doc, "7.2 Contacts & companies")
    bullets(doc, [
        "contacts[] — master contact DB; handlers/contacts.js",
        "companies[] + companies-hub.js — account aggregation (Phase 2 promotion)",
        "pipelineContact.js syncs lead fields ↔ master contact",
    ])
    h2(doc, "7.3 Deals & opportunities")
    bullets(doc, [
        "Primary: crm.deals[] nested per lead",
        "API: crm-deals.js, opportunities-hub.js",
        "Freight variant: FREIGHT_DEAL_STAGES for logistics orgs",
        "Actions: addDeal, closeDealWon, closeDealLost (crmWorkflow.js)",
    ])

    # 8
    h1(doc, "8. Application Shell & Navigation")
    h2(doc, "8.1 Shell layout")
    para(doc, (
        "AppShell.jsx orchestrates: Sidebar (navConfig.js), AppHeader, PanelViewport, "
        "LeadWorkspace overlay, CommandPalette, ConnectAssistant, onboarding modals."
    ))
    para(doc, "URL model: /home/dashboard?panel=pipeline&status=new&lead=<id>")
    h2(doc, "8.2 Navigation IA (customer)")
    bullets(doc, [
        "Home: overview, team intelligence (feature-gated), activity log (feature-gated)",
        "CRM/Sales: pipeline, contacts, opportunities, companies, marketing, calendar, automation",
        "Workspace: team members, permissions, hierarchy, email, WhatsApp settings",
        "Sidebar counts from pipeline bootstrap SQL/API — not client-side full scans",
    ])
    h2(doc, "8.3 Lead workspace (record pattern)")
    bullets(doc, [
        "LeadWorkspace.jsx — tabs: Overview, Deals, Timeline, Tasks & meetings, Email, WhatsApp",
        "Opened via openPipelineLead(leadId) without leaving shell",
        "Timeline: buildUnifiedTimeline() in crmTimeline.js merges all event types",
        "Return navigation from team intelligence via TeamIntelReturnBanner",
    ])

    # 9
    h1(doc, "9. Key User Flows")
    h2(doc, "9.1 Pipeline bootstrap & load")
    bullets(doc, [
        "1. GET /api/pipeline/bootstrap — counts, summary, generation token",
        "2. List/board/deals via pipelineListLoad.js — paginated, scoped",
        "3. Sidebar reflects pipelineSummary.byStatus",
    ])
    h2(doc, "9.2 Pipeline CRUD")
    table(doc, ["Action", "API / code", "Side effects"], [
        ("List/filter", "GET saved-leads.js", "Scope by role; SQL or shard"),
        ("Create", "POST saved-leads, extension capture", "Contact upsert, workflow rules"),
        ("Update", "PATCH saved-leads/:id", "Score, Meili sync, SQL sync, notifications"),
        ("Delete", "DELETE (permission-gated)", "Audit, index removal"),
        ("Bulk", "crm-bulk-update.js", "Batch workflow + sync"),
    ])
    h2(doc, "9.3 Team admin flow")
    bullets(doc, [
        "Org admin invites via inviteTeamMember() → email + organizationInvites",
        "Set pipelineRole, marketingRole, canSearch per member",
        "Team → Permissions tab edits role_permissions matrix",
        "Departments/teams: org/teams, org/departments handlers",
    ])

    # 10
    h1(doc, "10. CRM Activity, Team Metrics & Attribution")
    h2(doc, "10.1 Per-record timeline")
    para(doc, (
        "crmTimeline.js (frontend) and crm-activity-timeline.js (API) merge CRM activities, emails, "
        "marketing events, tasks, meetings, notes into one chronological feed per lead."
    ))
    h2(doc, "10.2 Org activity log")
    bullets(doc, [
        "Panel: crm-log (CrmActivityLogPanel) — feature flag ACTIVITY_LOG_HUB_IN_CRM_ENABLED",
        "API: crm-activity-log.js → readActivityLogCached, activityLogQuery.js",
        "Filters: period, member, activity type, tags, pagination",
    ])
    h2(doc, "10.3 Team intelligence & rep review")
    bullets(doc, [
        "APIs: crm-team-metrics, crm-team-dashboard, crm-dashboard-kpi, crm-rep-review",
        "buildRepReviewPayload() — single source for rep drill-down rollup + feed",
        "Period vocabulary: canonicalActivityPeriod — day | 7d | 30d (never mix calendar week with rolling 7d)",
    ])
    h2(doc, "10.4 Last active attribution (critical rule)")
    table(doc, ["Metric", "Source", "Must NOT use"], [
        ("Last active (CRM)", "resolveLastCrmActivityMap — actor_id per user", "Assignee only, login, workspace pulse"),
        ("In-app time", "aggregateWorkspaceUsage().lastInAppAt", "CRM activity timestamps"),
        ("Rep metrics scope", "Same since/until + actor filter for rollup and feed", "Lead assignee-only scans"),
    ])

    # 11
    h1(doc, "11. AI Copilot")
    bullets(doc, [
        "UI: ConnectAssistant.jsx — slide-over; context: panel, tab, leadId",
        "API: POST /api/assistant/chat → processCopilotTurn (lib/server/copilot/orchestrator.js)",
        "Pipeline: intent → planner → retrievers (CRM, Meili, web) → agents → synthesis → actions",
        "CRM tools: pipeline queries, follow-ups, stalled deals (crmTools.js)",
        "Email/WhatsApp generation: crm-generate-email, crm/generate-whatsapp",
        "Lead scoring: computeCrmLeadScore — org-configurable rules",
    ])

    # 12
    h1(doc, "12. Marketing Hub")
    bullets(doc, [
        "Panel: MarketingPanel — tabs: overview, campaigns, templates, automations, forms, audiences, analytics, domains",
        "Store: marketingCampaigns, lists, segments, enrollments, templates, forms, landingPages, automations, events",
        "Send path: Resend/SES; USE_MARKETING_SQL_QUEUE for Postgres queue at scale",
        "CRM integration: marketing events on lead timeline; bulk enroll from pipeline",
        "Permissions: access_marketing, send_campaigns in RBAC matrix + marketingRoles.js",
        "Cron: marketing/cron for scheduled sends; webhooks for opens/clicks",
    ])

    # 13
    h1(doc, "13. Search")
    bullets(doc, [
        "Platform search: GET /api/platform/search → searchPlatformFast",
        "Fallback chain: Meilisearch → pipeline SQL table search → JSON shard",
        "Types: lead, contact, company, deal, campaign, task, note",
        "RBAC: hits re-checked via isPipelineEntryVisibleAsync",
        "Sync: syncMeilisearchAfterSave on pipeline save; meili-sync-cron daily",
        "AI prospect search: search-leads API (nav disabled in CRM shell via feature flag)",
    ])

    # 14
    h1(doc, "14. Automations & Workflows")
    h2(doc, "14.1 CRM workflow rules")
    bullets(doc, [
        "Triggers: status_enter, lead_created, no_activity_days (crmWorkflowRules.js)",
        "Engine: workflowEngine.js → applyCrmWorkflowRules on save",
        "UI: CrmAutomationPanel + AutomationCanvas (org admin)",
    ])
    h2(doc, "14.2 Marketing automations")
    bullets(doc, [
        "Graph runner: automationGraphRunner.js",
        "Triggers: contact_added, email opened/clicked, form submitted",
        "Storage: marketingAutomations[], marketingAutomationRuns[]",
    ])
    h2(doc, "14.3 Sequences")
    bullets(doc, [
        "crmSequences, crmSequenceEnrollments collections",
        "API: crm/sequences; panel crm-sequences under Automation nav",
    ])

    # 15
    h1(doc, "15. Email Strategy")
    para(doc, (
        "Extension-first Gmail strategy (CASA deferred): Chrome extension for send-and-log and thread sync. "
        "Web OAuth gmail.send not required for customers at early stage."
    ))
    bullets(doc, [
        "Extension APIs: extension/bootstrap, lead-match, capture-lead, log, crm/sync-email-thread",
        "In-app: crm/send-email (Resend or connected Gmail), crm-generate-email (AI)",
        "Inbound: crm/email-inbound webhook; reply routing",
        "Org domain: org/email-domain — Resend verification; orgOutboundEmailReady on session",
        "Consent: commercialEmailOptIn on leads (emailConsent.js)",
    ])

    # 16
    h1(doc, "16. API Handler Pattern")
    para(doc, "Every handler follows: applyCors → handleOptions → requireUser/requireAdmin → business logic → sendJson.")
    table(doc, ["Pattern", "File", "Purpose"], [
        ("Route registry", "apiRouteRegistry.js", "Path → dynamic import handler"),
        ("Auth", "auth.js", "getSessionUser, requireUser, requireAdmin, requireOrgAdmin"),
        ("Store RMW", "store.js", "readStore, updateStorePartial, withStoreLock"),
        ("Permissions", "permissionEnforce.js", "assertOrgPermission, mapUserToPermissionRole"),
        ("Pipeline mutations", "handlers/saved-leads.js", "Hub save chain with sync side effects"),
    ])
    h3(doc, "16.1 Post-save side effect chain (pipeline)")
    bullets(doc, [
        "1. Persist lead (shard or table)",
        "2. Master contact upsert",
        "3. Workflow rules + marketing automations bridge",
        "4. Lead score recompute",
        "5. SQL sync (deals, tasks, meetings, companies, activities)",
        "6. Meilisearch sync",
        "7. Dashboard snapshot bump",
        "8. Assignment notifications",
    ])

    # 17
    h1(doc, "17. Data Storage Deep Dive")
    h2(doc, "17.1 JSON store collections (40+)")
    para(doc, (
        "Key collections: users, organizations, organizationMemberships, savedLeads, contacts, companies, "
        "marketingCampaigns, crmSequences, supportTickets, adminAuditLog. "
        "Production requires Supabase; local dev can use SQLite data/connect-intel.sqlite."
    ))
    h2(doc, "17.2 Performance slices")
    table(doc, ["Slice", "Collections", "Use"], [
        ("AUTH_STORE_COLLECTIONS", "users, orgs, memberships, invites", "Login/session"),
        ("INVITE_STORE_COLLECTIONS", "+ platform data", "Team invite"),
        ("META_STORE_COLLECTIONS", "users, orgs, memberships", "Visibility checks"),
    ])
    h2(doc, "17.3 Pipeline sharding")
    bullets(doc, [
        "Large orgs: pipelineShard.js splits JSON blobs",
        "Enterprise: pipeline_leads PostgreSQL table with indexed owner_id, team_id, status",
        "Bootstrap generation tokens invalidate client cache on bulk changes",
    ])

    # 18
    h1(doc, "18. Security & Compliance")
    bullets(doc, [
        "Session JWT with SESSION_SECRET required in production",
        "Platform admin: ADMIN_EMAILS + three-layer UI/API guards",
        "Tenant isolation on every handler write",
        "Audit: recordPipelineAudit, org audit log, platform adminAuditLog",
        "docs/SECURITY.md — TOTP for platform admins (roadmap/enforced per env)",
    ])

    # 19
    h1(doc, "19. Production & Deployment")
    table(doc, ["Step", "Command / action"], [
        ("Pre-ship", "npm run prod:ship — build + missing file checks"),
        ("Deploy", "git push origin main → GitHub CI → Vercel"),
        ("Post-deploy", "npm run prod:log — PRODUCTION_LOG.md + Meili sync"),
        ("Rollback", "npm run prod:rollback -- <commit>"),
        ("Ops", "npm run prod:ops — Meilisearch sync all orgs"),
    ])
    h3(doc, "19.1 Scheduled crons (vercel.json)")
    bullets(doc, [
        "*/5 min — grafana/metrics-cron",
        "04:00 UTC — crm/dashboard-warm-cron",
        "04:30 UTC — crm/meili-sync-cron",
        "09:00 UTC — reports/scheduled-cron",
    ])

    # 20
    h1(doc, "20. Codebase Map (Key Files)")
    table(doc, ["Area", "Path"], [
        ("SPA entry", "frontend/src/App.jsx"),
        ("App shell", "frontend/src/components/layout/AppShell.jsx"),
        ("Navigation", "frontend/src/lib/navConfig.js"),
        ("URL state", "frontend/src/lib/appHistory.js"),
        ("Pipeline UI", "frontend/src/components/crm/PipelinePanel.jsx"),
        ("Lead record", "frontend/src/components/crm/LeadWorkspace.jsx"),
        ("API router", "api/index.js"),
        ("Route registry", "lib/server/apiRouteRegistry.js"),
        ("Tenancy", "lib/server/organizations.js"),
        ("Pipeline save", "lib/server/handlers/saved-leads.js"),
        ("CRM core", "lib/server/crm.js, crmWorkflow.js"),
        ("Copilot", "lib/server/copilot/orchestrator.js"),
        ("Blueprint", "docs/CRM_PLATFORM_BLUEPRINT.md"),
    ])

    # 21
    h1(doc, "21. End-to-End Flow Diagrams (Narrative)")
    h2(doc, "21.1 New company user journey")
    bullets(doc, [
        "Landing → Auth (Google) → upsertUser → OnboardingModal (company name + domain check)",
        "completeOnboarding → org + membership (org_admin) → SQL sync deferred",
        "AppShell loads overview or pipeline → pipeline/bootstrap → PipelinePanel",
        "Create/import lead → saved-leads POST → contact sync → workflows → Meili → sidebar counts update",
        "Open lead → LeadWorkspace → timeline, tasks, email via extension or Resend",
    ])
    h2(doc, "21.2 Manager reviews rep performance")
    bullets(doc, [
        "Navigate to team intelligence (crm-dashboard) when feature enabled",
        "crm-team-metrics with canonicalActivityPeriod (7d or 30d)",
        "Drill rep → crm-rep-review → buildRepReviewPayload with shared since/until",
        "Last active column uses resolveLastCrmActivityMap (actor_id), not assignee",
    ])
    h2(doc, "21.3 Marketing campaign to CRM attribution")
    bullets(doc, [
        "Build audience in MarketingPanel → create campaign → enqueue sends (SQL queue)",
        "Opens/clicks → marketingEvents + marketing/webhooks",
        "Events appear on lead timeline (crmTimeline.js)",
        "Lead score may react to engagement (org scoring rules)",
    ])

    h1(doc, "Appendix A — Environment Flags")
    table(doc, ["Flag", "Effect"], [
        ("USE_PIPELINE_LEADS_TABLE", "SQL pipeline reads"),
        ("USE_PIPELINE_HIERARCHY_RBAC", "SQL role scoping"),
        ("USE_MARKETING_SQL_QUEUE", "Postgres email queue"),
        ("ADMIN_EMAILS", "Platform operator emails (comma-separated)"),
        ("MEILI_*", "Meilisearch connection"),
        ("SUPABASE_URL + SERVICE_ROLE_KEY", "Required in production"),
    ])

    h1(doc, "Appendix B — Related Documentation")
    bullets(doc, [
        "docs/CRM_PLATFORM_BLUEPRINT.md — product evolution master",
        "docs/API.md — API route groups",
        "docs/WORKFLOW_ENGINE.md — automation details",
        "docs/CONNECT_COPILOT_ARCHITECTURE.md — AI copilot",
        "docs/EMAIL_STRATEGY_EARLY_STAGE.md — Gmail/extension strategy",
        "docs/RELEASE_CHECKLIST.md — manual smoke tests",
        "docs/PRODUCTION_LOG.md — live commit history",
    ])

    doc.add_paragraph()
    p = doc.add_paragraph("— End of document —")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def main():
    doc = Document()
    set_doc_defaults(doc)
    build(doc)
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    main()
